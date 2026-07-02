"""
ingest.py — turn a vendor's proposal files into plain text the engine can score against.

Vendor proposals arrive (due July 2) as a mix of PDF, Word, Excel, and plain text.
This module extracts text from each, with graceful fallbacks if an optional parser
library is not installed. It also offers a light keyword-retrieval helper so the
scoring engine can pull the most relevant passages for a batch of requirements
instead of stuffing an entire 200-page proposal into every prompt.
"""
from __future__ import annotations
import os
import re
from typing import List, Optional, Tuple


def extract_text(path: str) -> str:
    """Dispatch on file extension. Returns extracted text (possibly empty)."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext in (".txt", ".md"):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        if ext == ".pdf":
            return _pdf(path)
        if ext in (".docx",):
            return _docx(path)
        if ext in (".xlsx", ".xlsm"):
            return _xlsx(path)
    except ImportError as e:
        return (f"[parser not installed for {os.path.basename(path)} ({ext}); "
                f"install the matching library (pdfplumber/pypdf, python-docx, or openpyxl): {e}]")
    except Exception as e:
        return f"[ingest error for {os.path.basename(path)}: {type(e).__name__}: {e}]"
    return f"[unsupported file type: {ext}]"


def extract_many(paths: List[str]) -> str:
    """Concatenate text from several files with clear separators."""
    chunks = []
    for p in paths:
        chunks.append(f"\n\n===== FILE: {os.path.basename(p)} =====\n")
        chunks.append(extract_text(p))
    return "".join(chunks)


def fetch_url(url: str, timeout: int = 30) -> str:
    """
    Download a URL and extract its text. Handles documents served over HTTP
    (PDF/DOCX/XLSX/TXT/MD) by saving to a temp file and reusing extract_text(), and
    HTML pages by stripping tags. Used when a vendor posts its response on a portal
    or shared link rather than (or in addition to) attaching files.

    Note: this is the LOCALLY-RUN APP fetching a URL the operator explicitly supplied —
    it is the tool's own functionality, intended for vendor proposal links.
    """
    import urllib.request
    import tempfile
    import html
    from urllib.parse import urlparse

    if not url.lower().startswith(("http://", "https://")):
        return f"[invalid url (must start with http/https): {url}]"

    # SSRF guard: refuse private/loopback/link-local/reserved addresses
    import ipaddress, socket
    host = urlparse(url).hostname or ""
    try:
        infos = socket.getaddrinfo(host, None)
        for fam, _, _, _, sockaddr in infos:
            ip = ipaddress.ip_address(sockaddr[0])
            if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved:
                return f"[refused: {url} resolves to a non-public address ({ip})]"
    except Exception as e:
        return f"[refused: could not resolve host for {url}: {type(e).__name__}: {e}]"

    try:
        req = urllib.request.Request(url, headers={"User-Agent": "FSM-RFP-Eval-Agent/1.0"})
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            ctype = (resp.headers.get("Content-Type") or "").lower()
            max_bytes = int(os.environ.get("MAX_UPLOAD_MB", "25")) * 1024 * 1024
            data = resp.read(max_bytes + 1)
            if len(data) > max_bytes:
                return f"[refused: {url} exceeds the {max_bytes // (1024*1024)} MB limit]"
    except Exception as e:
        return f"[url fetch error for {url}: {type(e).__name__}: {e}]"

    path_ext = os.path.splitext(urlparse(url).path)[1].lower()
    doc_exts = (".pdf", ".docx", ".xlsx", ".xlsm", ".txt", ".md")

    # Document by extension or content-type -> save to temp and reuse extract_text().
    is_doc = path_ext in doc_exts or any(
        k in ctype for k in ("pdf", "officedocument", "spreadsheet", "msword"))
    if is_doc:
        ext = path_ext if path_ext in doc_exts else (
            ".pdf" if "pdf" in ctype else
            ".xlsx" if "spreadsheet" in ctype else
            ".docx" if ("officedocument" in ctype or "msword" in ctype) else ".txt")
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tf:
            tf.write(data)
            tmp = tf.name
        try:
            return extract_text(tmp)
        finally:
            try:
                os.unlink(tmp)
            except OSError:
                pass

    # Otherwise treat as text/HTML: decode and strip tags.
    text = data.decode("utf-8", errors="ignore")
    if "html" in ctype or "<html" in text[:2000].lower() or "<body" in text[:4000].lower():
        text = re.sub(r"(?is)<(script|style|head|nav|footer)[^>]*>.*?</\1>", " ", text)
        text = re.sub(r"(?s)<[^>]+>", " ", text)        # strip remaining tags
        text = html.unescape(text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    return text.strip()


def extract_sources(paths: Optional[List[str]] = None,
                    urls: Optional[List[str]] = None) -> str:
    """
    Build a single proposal-text blob from any mix of uploaded files and URLs,
    each clearly delimited so retrieval can attribute passages back to a source.
    """
    parts = []
    for p in (paths or []):
        parts.append(f"\n\n===== FILE: {os.path.basename(p)} =====\n")
        parts.append(extract_text(p))
    for u in (urls or []):
        u = u.strip()
        if not u:
            continue
        parts.append(f"\n\n===== URL: {u} =====\n")
        parts.append(fetch_url(u))
    return "".join(parts).strip()


def _pdf(path: str) -> str:
    try:
        import pdfplumber  # preferred: better layout handling
        out = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                out.append(f"===== LOC: p.{i + 1} =====")
                out.append(page.extract_text() or "")
        return "\n".join(out)
    except ImportError:
        from pypdf import PdfReader  # fallback
        reader = PdfReader(path)
        out = []
        for i, pg in enumerate(reader.pages):
            out.append(f"===== LOC: p.{i + 1} =====")
            out.append(pg.extract_text() or "")
        return "\n".join(out)


def _docx(path: str) -> str:
    from docx import Document  # python-docx
    parts = []
    doc = Document(path)
    for p in doc.paragraphs:
        style = (p.style.name if p.style else "") or ""
        if style.lower().startswith("heading") and p.text.strip():
            parts.append(f"===== LOC: {p.text.strip()} =====")
        parts.append(p.text)
    for table in doc.tables:                     # vendor responses often live in tables
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _xlsx(path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"===== LOC: Sheet '{ws.title}' =====")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


# --------------------------------------------------------------------------- #
# Lightweight retrieval (no embeddings needed)                                #
# --------------------------------------------------------------------------- #
def chunk_text(text: str, target_chars: int = 1200) -> List[str]:
    """Split text into ~target_chars chunks on paragraph boundaries."""
    paras = re.split(r"\n\s*\n", text)
    chunks, cur = [], ""
    for para in paras:
        if len(cur) + len(para) > target_chars and cur:
            chunks.append(cur.strip())
            cur = ""
        cur += para + "\n\n"
    if cur.strip():
        chunks.append(cur.strip())
    return chunks


RetrievalIndex = List[Tuple[str, str]]  # (original chunk, lowercase chunk)


def build_retrieval_index(text: str, target_chars: int = 1200) -> RetrievalIndex:
    """Pre-chunk/lowercase proposal text so repeated retrieval calls stay cheap."""
    return [(chunk, chunk.lower()) for chunk in chunk_text(text, target_chars)]


def relevant_passages(
    text: str,
    keywords: List[str],
    max_chunks: int = 6,
    index: Optional[RetrievalIndex] = None,
) -> str:
    """
    Return the chunks most relevant to a set of keywords (e.g. requirement terms),
    by simple term-overlap scoring. Good enough to localize a vendor's answer without
    a vector store; the engine can pass requirement keywords for a scoring batch.
    """
    entries = index if index is not None else build_retrieval_index(text)
    kw = [k.lower() for k in keywords if len(k) > 3]
    if not kw:
        return "\n\n".join(chunk for chunk, _ in entries[:max_chunks])
    scored = []
    for ch, low in entries:
        score = sum(low.count(k) for k in kw)
        if score:
            scored.append((score, ch))
    scored.sort(key=lambda x: x[0], reverse=True)
    top = [c for _, c in scored[:max_chunks]]
    return "\n\n".join(top) if top else "\n\n".join(chunk for chunk, _ in entries[:max_chunks])


# --------------------------------------------------------------------------- #
# Requirements-matrix alignment (join vendor responses to RIDs)               #
# --------------------------------------------------------------------------- #
_RID_CELL_RE = re.compile(r"^[A-Z]{2,5}-\d{1,4}$")


def _cell(row, idx) -> str:
    if idx is None or idx >= len(row):
        return ""
    v = row[idx]
    return "" if v is None else str(v).strip()


def _find_rid_column(rows, known_rids):
    """(header_row_index, rid_col_index) or (None, None). Prefer a 'Req ID'/'RID'
    header; else the column whose values best match known RIDs."""
    for hi, row in enumerate(rows[:6]):
        for ci, cell in enumerate(row):
            name = re.sub(r"\s+", " ", str(cell or "")).strip().lower().rstrip(".")
            if name in ("req id", "rid", "requirement id", "req id#", "req. id"):
                return hi, ci
    best_col, best_hits = None, 0
    ncols = max((len(r) for r in rows), default=0)
    for ci in range(ncols):
        hits = sum(1 for r in rows if _cell(r, ci).upper() in known_rids)
        if hits > best_hits:
            best_col, best_hits = ci, hits
    if best_col is not None and best_hits >= 3:
        for hi, r in enumerate(rows):
            if _cell(r, best_col).upper() in known_rids:
                # Walk upward from the row above the first RID hit, skipping
                # section-like rows (fewer than 2 non-empty cells), to find
                # the real header row instead of assuming it's immediately above.
                header_hi = max(0, hi - 1)
                while header_hi > 0:
                    non_empty = sum(1 for ci in range(len(rows[header_hi])) if _cell(rows[header_hi], ci))
                    if non_empty >= 2:
                        break
                    header_hi -= 1
                return header_hi, best_col
    return None, None


def _find_response_columns(rows, header_idx):
    """(code_col, response_col). Prefer headers containing 'response' (shortest avg
    cell = code, longest = narrative); else the last two data-bearing columns."""
    header = rows[header_idx] if 0 <= header_idx < len(rows) else ()
    resp_cols = [ci for ci, c in enumerate(header) if "response" in str(c or "").lower()]
    data = rows[header_idx + 1:]

    def avg_len(ci):
        vals = [len(_cell(r, ci)) for r in data if _cell(r, ci)]
        return sum(vals) / len(vals) if vals else 0.0

    if len(resp_cols) >= 2:
        resp_cols.sort(key=avg_len)
        return resp_cols[0], resp_cols[-1]
    if len(resp_cols) == 1:
        return resp_cols[0], resp_cols[0]
    ncols = max((len(r) for r in rows), default=0)
    filled = [ci for ci in range(ncols) if any(_cell(r, ci) for r in data)]
    if len(filled) >= 2:
        return filled[-2], filled[-1]
    return (filled[-1], filled[-1]) if filled else (None, None)


def _find_requirement_text_column(rows, header_idx, text_index):
    """Return the column index whose cell values best match known requirement
    texts (normalized), or None. Used only when no RID column is found."""
    data = rows[header_idx + 1:]
    ncols = max((len(r) for r in rows), default=0)
    best_col, best_hits = None, 0
    for ci in range(ncols):
        hits = sum(1 for r in data if _norm(_cell(r, ci)) in text_index)
        if hits > best_hits:
            best_col, best_hits = ci, hits
    return best_col if best_hits >= 3 else None


def extract_requirement_matrix(paths, requirements):
    """Parse a submitted requirements matrix (.xlsx/.xlsm) into {rid: {code,
    response, source, sheet}} by joining on the RID column. Returns {} when no
    matrix is present or openpyxl is unavailable — callers then behave as before."""
    known = {str(r["rid"]).strip().upper() for r in (requirements or [])}
    out: dict = {}
    if not known:
        return out
    for p in (paths or []):
        if os.path.splitext(p)[1].lower() not in (".xlsx", ".xlsm"):
            continue
        try:
            import openpyxl
            wb = openpyxl.load_workbook(p, read_only=True, data_only=True)
        except Exception:
            continue
        fname = os.path.basename(p)
        try:
            for ws in wb.worksheets:
                rows = [tuple(r) for r in ws.iter_rows(values_only=True)]
                if not rows:
                    continue
                hi, rid_col = _find_rid_column(rows, known)
                text_index = {_norm(r["requirement"]): str(r["rid"]).strip().upper()
                              for r in requirements if r.get("requirement")}
                text_col = None
                if rid_col is None:
                    # header row for a text-only sheet is row 0 unless a match run starts lower
                    hi = 0
                    text_col = _find_requirement_text_column(rows, hi, text_index)
                    if text_col is None:
                        continue
                code_col, resp_col = _find_response_columns(rows, hi)
                for row in rows[hi + 1:]:
                    if rid_col is not None:
                        rid = _cell(row, rid_col).upper()
                    else:
                        rid = text_index.get(_norm(_cell(row, text_col)), "")
                    if rid not in known or rid in out:
                        continue
                    code = _cell(row, code_col)
                    resp = _cell(row, resp_col)
                    if code_col == resp_col:
                        code = ""
                    if not (code or resp):
                        continue
                    out[rid] = {"code": code, "response": resp, "source": fname, "sheet": ws.title}
        finally:
            try:
                wb.close()
            except Exception:
                pass
    return out


# --------------------------------------------------------------------------- #
# Source segmentation & evidence locators                                     #
# --------------------------------------------------------------------------- #
_SRC_RE = re.compile(r"^=====\s+(?:FILE|URL):\s+(.*?)\s+=====$")
_LOC_RE = re.compile(r"^=====\s+LOC:\s+(.*?)\s+=====$")


def _norm(s: str) -> str:
    """Lowercase, collapse whitespace, strip surrounding quotes — for tolerant
    quote-to-segment matching."""
    s = re.sub(r"\s+", " ", (s or "")).strip()
    return s.strip("\"'“”‘’").lower()  # ASCII + curly quotes


def parse_segments(blob: str) -> List[dict]:
    """Reconstruct (source, locator, text) segments from a blob carrying
    ===== FILE/URL: ... ===== source headers and ===== LOC: ... ===== markers.
    A blob with no headers yields a single (source='', locator='(document)')
    segment. Each segment also caches its normalized text under '_norm' for fast
    locate_quote()."""
    segments: List[dict] = []
    source, locator = "", "(document)"
    buf: List[str] = []

    def flush():
        text = "\n".join(buf).strip()
        if text:
            segments.append({"source": source, "locator": locator,
                             "text": text, "_norm": _norm(text)})
        buf.clear()

    for line in (blob or "").split("\n"):
        m = _SRC_RE.match(line)
        if m:
            flush()
            source = m.group(1)
            locator = "(document)"      # reset locator when the source changes
            continue
        m = _LOC_RE.match(line)
        if m:
            flush()
            locator = m.group(1)
            continue
        buf.append(line)
    flush()
    return segments


def strip_loc_markers(blob: str) -> str:
    """Remove only ===== LOC: ... ===== lines, leaving FILE/URL headers and text.
    On the text/markdown/PDF/URL/sample paths this reproduces the pre-evidence
    blob byte-for-byte, so scoring is provably unperturbed (the regression
    anchor). XLSX differs only by the intentional '## Sheet:' -> LOC swap."""
    return "\n".join(line for line in (blob or "").split("\n")
                     if not _LOC_RE.match(line))


def locate_quote(quote: str, segments: List[dict], min_len: int = 12) -> dict:
    """Map a verbatim quote back to the segment that contains it (normalized
    substring). Returns {source, locator}; {source:'', locator:'(unlocated)'}
    when the quote is too short or not found — never fabricates a locator."""
    nq = _norm(quote)
    if len(nq) < min_len:
        return {"source": "", "locator": "(unlocated)"}
    for seg in segments:
        seg_norm = seg.get("_norm")
        if seg_norm is None:
            seg_norm = _norm(seg["text"])
        if nq in seg_norm:
            return {"source": seg["source"], "locator": seg["locator"]}
    return {"source": "", "locator": "(unlocated)"}
